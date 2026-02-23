#!/usr/bin/env python3
"""
Universal Unit Import Tool
==========================
Parses a Monograph Word doc, resolves templates, imports defects.

Usage:
  python3 universal_import.py <docx_path> <cycle_id>            # dry run
  python3 universal_import.py <docx_path> <cycle_id> --commit    # actual import

Requires: python-docx (pip install python-docx)
"""
import sqlite3
import uuid
import sys
import os
from datetime import datetime, timezone
from difflib import SequenceMatcher

DB_PATH = '/var/data/inspections.db'
TENANT = 'MONOGRAPH'

# ============================================================
# 1. AREA HEADER MAP (Bold text in Word doc -> DB area name)
# ============================================================
AREA_HEADERS = {
    'kitchen': 'KITCHEN',
    'lounge area': 'LOUNGE',
    'lounge': 'LOUNGE',
    'bedroom a': 'BEDROOM A',
    'bedroom b': 'BEDROOM B',
    'bedroom c': 'BEDROOM C',
    'bedroom d': 'BEDROOM D',
    'bathroom': 'BATHROOM',
}

# These headers mean "stop parsing defects"
STOP_SECTIONS = {'general notes', 'general', 'items that will not be tested',
                 'items not tested', 'notes'}

# ============================================================
# 2. SUB-HEADER MAP: (area, sub_header_lower) -> (category, parent_kw)
#    parent_kw=None means defects match directly against category children
# ============================================================
SUB_HEADER_MAP = {
    # Universal
    ('*', 'walls'): ('WALLS', None),
    ('*', 'wall'): ('WALLS', None),
    ('*', 'floor'): ('FLOOR', None),
    ('*', 'ceiling'): ('CEILING', None),
    ('*', 'electrical'): ('ELECTRICAL', None),
    ('*', 'frame'): ('DOORS', 'Frame'),
    ('*', 'ironmongery'): ('DOORS', 'Ironmongery'),

    # Kitchen-specific
    ('KITCHEN', 'door'): ('DOORS', 'D1'),
    ('KITCHEN', 'doors'): ('DOORS', 'D1'),
    ('KITCHEN', 'window'): ('WINDOWS', 'W1'),
    ('KITCHEN', 'windows'): ('WINDOWS', 'W1'),
    ('KITCHEN', 'sink pack'): ('JOINERY', 'Sink pack'),
    ('KITCHEN', 'bin drawer'): ('JOINERY', 'Bin drawer'),
    ('KITCHEN', 'drawer pack'): ('JOINERY', 'Drawer pack'),
    ('KITCHEN', 'counter seating'): ('JOINERY', 'Counter seating'),
    ('KITCHEN', 'lockable pack 1&2'): ('JOINERY', 'Lockable pack 1 & 2'),
    ('KITCHEN', 'lockable pack 1 & 2'): ('JOINERY', 'Lockable pack 1 & 2'),
    ('KITCHEN', 'lockable pack 3&4'): ('JOINERY', 'Lockable pack 3 & 4'),
    ('KITCHEN', 'lockable pack 3 & 4'): ('JOINERY', 'Lockable pack 3 & 4'),
    ('KITCHEN', 'eye level pack'): ('JOINERY', 'Eye level pack'),

    # Bedroom-specific
    ('BEDROOM A', 'door'): ('DOORS', 'D2'),
    ('BEDROOM A', 'doors'): ('DOORS', 'D2'),
    ('BEDROOM A', 'window'): ('WINDOWS', 'W3'),
    ('BEDROOM A', 'windows'): ('WINDOWS', 'W3'),
    ('BEDROOM B', 'door'): ('DOORS', 'D2'),
    ('BEDROOM B', 'doors'): ('DOORS', 'D2'),
    ('BEDROOM B', 'window'): ('WINDOWS', 'W3'),
    ('BEDROOM B', 'windows'): ('WINDOWS', 'W3'),
    ('BEDROOM C', 'door'): ('DOORS', 'D2'),
    ('BEDROOM C', 'doors'): ('DOORS', 'D2'),
    ('BEDROOM C', 'window'): ('WINDOWS', 'W3'),
    ('BEDROOM C', 'windows'): ('WINDOWS', 'W3'),
    ('BEDROOM D', 'door'): ('DOORS', 'D2'),
    ('BEDROOM D', 'doors'): ('DOORS', 'D2'),
    ('BEDROOM D', 'window'): ('WINDOWS', 'W3'),
    ('BEDROOM D', 'windows'): ('WINDOWS', 'W3'),

    # Bedroom joinery
    ('*', 'b.i.c'): ('JOINERY', 'B.I.C.'),
    ('*', 'b.i.c.'): ('JOINERY', 'B.I.C.'),
    ('*', 'bic'): ('JOINERY', 'B.I.C.'),
    ('*', 'floating shelf'): ('JOINERY', 'Floating shelf'),
    ('*', 'study desk'): ('JOINERY', 'Study desk'),

    # Bathroom-specific
    ('BATHROOM', 'door'): ('DOORS', 'D2a'),
    ('BATHROOM', 'doors'): ('DOORS', 'D2a'),
    ('BATHROOM', 'window'): ('WINDOWS', 'W2'),
    ('BATHROOM', 'windows'): ('WINDOWS', 'W2'),
    ('BATHROOM', 'd3'): ('DOORS', 'D3'),
    ('BATHROOM', 'wc'): ('PLUMBING', 'WC'),
    ('BATHROOM', 'mixer'): ('PLUMBING', 'Mixer'),
    ('BATHROOM', 'arm'): ('PLUMBING', 'Arm'),
    ('BATHROOM', 'shut off valve cold'): ('PLUMBING', 'Shut off Cold'),
    ('BATHROOM', 'shut of valve cold'): ('PLUMBING', 'Shut off Cold'),
    ('BATHROOM', 'shut off cold'): ('PLUMBING', 'Shut off Cold'),
    ('BATHROOM', 'shut off valve hot'): ('PLUMBING', 'Shut off Hot'),
    ('BATHROOM', 'shut of valve hot'): ('PLUMBING', 'Shut off Hot'),
    ('BATHROOM', 'shut off hot'): ('PLUMBING', 'Shut off Hot'),
    ('BATHROOM', 'airbrick'): ('WALLS', 'Airbrick in shower'),
}

# Recognized sub-header strings (used to detect headers vs defect text)
KNOWN_SUB_HEADERS = set()
for (_, sh) in SUB_HEADER_MAP.keys():
    KNOWN_SUB_HEADERS.add(sh)

# ============================================================
# 3. STANDARD DROPS (auto-filtered)
# ============================================================
# (area, sub_header) combos to always drop
AREA_SUB_DROPS = {
    ('KITCHEN', 'frame'),   # Kitchen front door - always excluded
    ('KITCHEN', 'door'),
    ('KITCHEN', 'doors'),
}

# Content patterns to drop (case-insensitive substring match)
CONTENT_DROPS = [
    'wi-fi', 'wifi', 'wi fi',
    'plugs not tested',
    'no water in the',
    'no water in unit',
    'stove control panel',
    'shower control panel',
    'fire blanket',
    'cylinder lockset',
    'all ff&e', 'all ffe',
    'front door not inspected',
    'no hot water',
    'towel rail',
    'clothing rail',
    'blind', # window blinds FF&E
    'panel heater',  # physical unit is FF&E, only plug matters
    'not inspected as per',
]

# ============================================================
# 4. TEMPLATE OVERRIDES (verified against DB, consolidated from all imports)
# ============================================================
TEMPLATE_OVERRIDES = {
    # Kitchen root items
    ('KITCHEN', 'WALLS', 'paint'): '16e941da',
    ('KITCHEN', 'ELECTRICAL', 'DB'): '7414ad92',
    ('KITCHEN', 'FLOOR', 'Soft joint'): 'bdafda18',
    ('KITCHEN', 'CEILING', 'plaster recess'): '6eb3af36',

    # Lounge
    ('LOUNGE', 'WALLS', 'Wall', 'paint'): 'c248c406',
    ('LOUNGE', 'FLOOR', 'Floor', 'chipped'): '3cb1b144',
    ('LOUNGE', 'FLOOR', 'Floor', 'grout'): 'feafbe9d',
    ('LOUNGE', 'FLOOR', 'Floor', 'skirting'): 'a46f716d',
    ('LOUNGE', 'CEILING', 'plaster recess'): 'a4163cb8',
    ('LOUNGE', 'ELECTRICAL', 'Double plug'): 'fa47bce5',

    # Bedroom floors (all 12)
    ('BEDROOM A', 'FLOOR', 'Floor', 'chipped'): '41c9bd11',
    ('BEDROOM A', 'FLOOR', 'Floor', 'grout'): 'e1d9e932',
    ('BEDROOM A', 'FLOOR', 'Floor', 'skirting'): '14eb7511',
    ('BEDROOM B', 'FLOOR', 'Floor', 'chipped'): '2ed16ab7',
    ('BEDROOM B', 'FLOOR', 'Floor', 'grout'): '81520953',
    ('BEDROOM B', 'FLOOR', 'Floor', 'skirting'): '1136f030',
    ('BEDROOM C', 'FLOOR', 'Floor', 'chipped'): '85620ac4',
    ('BEDROOM C', 'FLOOR', 'Floor', 'grout'): 'f34b4fe9',
    ('BEDROOM C', 'FLOOR', 'Floor', 'skirting'): '6a0771ae',
    ('BEDROOM D', 'FLOOR', 'Floor', 'chipped'): '54ac6a45',
    ('BEDROOM D', 'FLOOR', 'Floor', 'grout'): '956b6837',
    ('BEDROOM D', 'FLOOR', 'Floor', 'skirting'): 'a39a8899',

    # Bedroom walls
    ('BEDROOM C', 'WALLS', 'Wall', 'finish'): '5628303a',

    # Bedroom study desk lights
    ('BEDROOM B', 'ELECTRICAL', 'study desk light'): 'e2fd6318',
    ('BEDROOM D', 'ELECTRICAL', 'Study desk light', 'screws'): 'b1b7e7ec',

    # Floating shelves (all 8)
    ('BEDROOM A', 'JOINERY', 'Floating shelf', 'finish'): '468ece9d',
    ('BEDROOM B', 'JOINERY', 'Floating shelf', 'finish'): '262bfbeb',
    ('BEDROOM C', 'JOINERY', 'Floating shelf', 'finish'): '2f006892',
    ('BEDROOM D', 'JOINERY', 'Floating shelf', 'finish'): '135828f3',
    ('BEDROOM A', 'JOINERY', 'Floating shelf', 'installed'): '519d4580',
    ('BEDROOM B', 'JOINERY', 'Floating shelf', 'installed'): '7b3e816b',
    ('BEDROOM C', 'JOINERY', 'Floating shelf', 'installed'): '0b929eb3',
    ('BEDROOM D', 'JOINERY', 'Floating shelf', 'installed'): 'f653cf83',

    # Bathroom
    ('BATHROOM', 'FLOOR', 'Floor', 'chipped'): '818a1716',
    ('BATHROOM', 'FLOOR', 'Floor', 'grout'): '8fa8781c',
    ('BATHROOM', 'CEILING', 'Ceiling', 'paint'): 'faea42ac',
    ('BATHROOM', 'PLUMBING', 'WC', 'installation'): '8667f32c',
    ('BATHROOM', 'PLUMBING', 'WC', 'shut off'): 'b9805e6c',
    ('BATHROOM', 'PLUMBING', 'Shut off Cold'): 'a9e99c5e',
    ('BATHROOM', 'PLUMBING', 'Shut off Hot'): 'e5372c9d',
    ('BATHROOM', 'PLUMBING', 'Arm'): '019d6605',

    # Bathroom walls
    ('BATHROOM', 'WALLS', 'tile trim', 'duct'): 'df84942f',
    ('BATHROOM', 'WALLS', 'tile', 'window sill'): '76c93f42',
    ('BATHROOM', 'WALLS', 'tile trim', 'window reveal'): '347c7f63',
    ('BATHROOM', 'WALLS', 'tile', 'broken'): 'ef937d8f',
    ('BATHROOM', 'FLOOR', 'mosaic', 'dark grey grout'): 'd5c6a122',
    ('BATHROOM', 'DOORS', 'D3', 'operation'): 'a6939da2',
    ('BATHROOM', 'DOORS', 'Ironmongery', 'handle'): '9d6fe4a5',

    # Bedroom D ironmongery
    ('BEDROOM D', 'DOORS', 'Ironmongery', 'handle'): '9d6fe4a5',
}

# Keywords used to search override dict against defect text
OVERRIDE_KEYWORDS = [
    'paint', 'orchid', 'chipped', 'grout', 'skirting', 'tile',
    'soft joint', 'plaster recess', 'crack',
    'DB', 'installation', 'operation', 'shut off',
    'duct', 'window sill', 'window reveal', 'broken',
    'dark grey', 'mosaic', 'handle', 'screws',
    'finish', 'installed', 'loose',
    'study desk light',
]

# ============================================================
# 5. HELPERS
# ============================================================
def gen_id():
    return uuid.uuid4().hex[:8]

def now_iso():
    return datetime.now(timezone.utc).isoformat()

def fuzzy(a, b):
    return SequenceMatcher(None, a.lower().strip(), b.lower().strip()).ratio()

def is_sub_header(text, current_area):
    """Check if text is a sub-header (not a defect description)."""
    t = text.lower().strip().rstrip('.')
    # Exact match against known sub-headers
    if t in KNOWN_SUB_HEADERS:
        return True
    # Check area-specific keys
    if (current_area, t) in SUB_HEADER_MAP:
        return True
    if ('*', t) in SUB_HEADER_MAP:
        return True
    # Fuzzy match for typos (e.g., "Shut of valve Hot")
    for sh in KNOWN_SUB_HEADERS:
        if len(t) < 30 and fuzzy(t, sh) >= 0.85:
            return True
    return False

def get_sub_header_mapping(area, sub_header_text):
    """Map sub-header to (category, parent_kw)."""
    t = sub_header_text.lower().strip().rstrip('.')
    # Area-specific first
    if (area, t) in SUB_HEADER_MAP:
        return SUB_HEADER_MAP[(area, t)]
    # Wildcard
    if ('*', t) in SUB_HEADER_MAP:
        return SUB_HEADER_MAP[('*', t)]
    # Fuzzy match
    best_key = None
    best_score = 0
    for key in SUB_HEADER_MAP:
        a, sh = key
        if a not in (area, '*'):
            continue
        s = fuzzy(t, sh)
        if s > best_score:
            best_score = s
            best_key = key
    if best_score >= 0.85 and best_key:
        return SUB_HEADER_MAP[best_key]
    return None, None

def detect_defect_type(text):
    """NI if clearly not installed/missing, else NTS."""
    t = text.lower()
    if 'not installed' in t and 'control panel' not in t:
        return 'NI'
    return 'NTS'

def should_drop_content(text):
    """Check if defect text matches a content-based drop pattern."""
    t = text.lower()
    for pattern in CONTENT_DROPS:
        if pattern.lower() in t:
            return True
    return False

# ============================================================
# 6. WORD DOC PARSER
# ============================================================
def parse_word_doc(path):
    """
    Parse Word doc -> (unit_number, inspector_name, inspector_date, defects)
    defects = [(area, sub_header, defect_text, line_num)]
    """
    from docx import Document
    doc = Document(path)

    # Extract metadata from tables
    unit_number = None
    inspector_name = None
    inspection_date = None

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            joined = ' '.join(cells).lower()
            # Unit number
            if 'unit no' in joined or 'area of inspection' in joined:
                for c in reversed(cells):
                    if c and c not in ('', 'UNIT NO / AREA OF INSPECTION'):
                        unit_number = c.strip()
                        break
            # Date
            if joined.startswith('date') or 'date' in cells[0].lower():
                for c in reversed(cells):
                    if c and c != cells[0]:
                        raw_date = c.strip()
                        # Convert DD.MM.YYYY to YYYY-MM-DD
                        try:
                            parts = raw_date.split('.')
                            if len(parts) == 3:
                                inspection_date = f"{parts[2]}-{parts[1]}-{parts[0]}"
                            else:
                                inspection_date = raw_date
                        except:
                            inspection_date = raw_date
                        break
            # Inspector
            if 'inspected by' in joined:
                for c in cells:
                    if 'inspected by' in c.lower():
                        name = c.replace('Inspected By:', '').replace('Inspected by:', '').strip()
                        if name:
                            inspector_name = name
                        break

    # Parse paragraphs for defects
    current_area = None
    current_sub_header = None
    defects = []
    parsing_active = True

    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue

        # Check if bold (area header)
        is_bold = any(r.bold for r in p.runs if r.text.strip()) if p.runs else False

        # Check for stop sections
        if text.lower().strip().rstrip(':') in STOP_SECTIONS:
            parsing_active = False
            continue

        if not parsing_active:
            continue

        # Area header (bold)
        if is_bold:
            area_key = text.lower().strip()
            if area_key in AREA_HEADERS:
                current_area = AREA_HEADERS[area_key]
                current_sub_header = None
                continue
            elif area_key in STOP_SECTIONS:
                parsing_active = False
                continue

        # Skip if no area set yet
        if not current_area:
            continue

        # Sub-header or defect?
        if is_sub_header(text, current_area):
            current_sub_header = text.strip()
            continue

        # It's a defect description
        if current_sub_header is None:
            # Defect with no sub-header (rare, skip)
            continue

        defects.append((current_area, current_sub_header, text, i))

    return unit_number, inspector_name, inspection_date, defects

# ============================================================
# 7. INSPECTOR LOOKUP
# ============================================================
INSPECTOR_MAP = {
    'alex nataniel': 'team-lead',
    'fisokuhle matsepe': 'insp-006',
    'fiso matsepe': 'insp-006',
    'lindokuhle zulu': 'insp-005',
    'stemi tumona': 'insp-001',
    'thebe majodina': 'insp-003',
    'thembinkosi biko': 'insp-004',
}

def get_inspector_id(name):
    if not name:
        return None
    key = name.lower().strip()
    if key in INSPECTOR_MAP:
        return INSPECTOR_MAP[key]
    # Fuzzy match
    for k, v in INSPECTOR_MAP.items():
        if fuzzy(key, k) >= 0.8:
            return v
    return None

# ============================================================
# 8. TEMPLATE RESOLVER
# ============================================================
def resolve_template(cur, area, category, parent_kw, defect_text):
    """
    Find the best matching template ID for a defect.
    Returns (template_id, match_info) or (None, fail_info).
    """
    dt = defect_text.lower()

    # --- Pass 1: Check overrides ---
    # Try exact 4-key matches with keywords from defect text
    for kw in OVERRIDE_KEYWORDS:
        if kw.lower() in dt or (parent_kw and kw.lower() in parent_kw.lower()):
            key4 = (area, category, parent_kw, kw)
            if key4 in TEMPLATE_OVERRIDES:
                return TEMPLATE_OVERRIDES[key4], f'OVERRIDE-4 {key4}'
            # Try case variations
            for ok in TEMPLATE_OVERRIDES:
                if len(ok) == 4 and ok[0] == area and ok[1] == category:
                    if (ok[2].lower() == (parent_kw or '').lower() and
                        ok[3].lower() == kw.lower()):
                        return TEMPLATE_OVERRIDES[ok], f'OVERRIDE-4 {ok}'

    # Try 3-key override with parent_kw
    key3 = (area, category, parent_kw)
    if key3 in TEMPLATE_OVERRIDES:
        return TEMPLATE_OVERRIDES[key3], f'OVERRIDE-3 {key3}'
    # Case-insensitive 3-key
    for ok in TEMPLATE_OVERRIDES:
        if len(ok) == 3 and ok[0] == area and ok[1] == category:
            if ok[2].lower() == (parent_kw or '').lower():
                return TEMPLATE_OVERRIDES[ok], f'OVERRIDE-3 {ok}'

    # Try 3-key with keywords from defect text
    for kw in OVERRIDE_KEYWORDS:
        if kw.lower() in dt:
            key3k = (area, category, kw)
            if key3k in TEMPLATE_OVERRIDES:
                return TEMPLATE_OVERRIDES[key3k], f'OVERRIDE-3kw {key3k}'
            for ok in TEMPLATE_OVERRIDES:
                if len(ok) == 3 and ok[0] == area and ok[1] == category:
                    if ok[2].lower() == kw.lower():
                        return TEMPLATE_OVERRIDES[ok], f'OVERRIDE-3kw {ok}'

    # --- Pass 2: Fuzzy match against DB templates ---
    cur.execute("""
        SELECT it.id, it.item_description, it.parent_item_id,
               pit.item_description as parent_desc
        FROM item_template it
        JOIN category_template ct ON it.category_id = ct.id
        JOIN area_template at2 ON ct.area_id = at2.id
        LEFT JOIN item_template pit ON it.parent_item_id = pit.id
        WHERE at2.area_name = ? AND ct.category_name = ? AND it.tenant_id = ?
        AND it.depth > 0
        ORDER BY it.item_order
    """, (area, category, TENANT))
    rows = cur.fetchall()

    best_id = None
    best_score = 0
    best_desc = ''
    best_parent = ''

    for tid, desc, pid, pdesc in rows:
        # Check parent match
        parent_match = False
        if parent_kw:
            if pdesc and parent_kw.lower() in pdesc.lower():
                parent_match = True
            elif parent_kw.lower() in desc.lower():
                parent_match = True
            elif pdesc and fuzzy(parent_kw, pdesc) >= 0.6:
                parent_match = True
        else:
            parent_match = True  # No parent filter

        if not parent_match:
            continue

        # Score: combine defect text match with item description match
        score = fuzzy(dt, desc)

        # Boost if key words from description appear in defect text
        desc_words = desc.lower().split()
        for w in desc_words:
            if len(w) > 3 and w in dt:
                score += 0.1

        if score > best_score:
            best_score = score
            best_id = tid
            best_desc = desc
            best_parent = pdesc or ''

    # Broader search if parent-filtered search failed
    if best_score < 0.3 and parent_kw:
        for tid, desc, pid, pdesc in rows:
            score = fuzzy(dt, desc)
            desc_words = desc.lower().split()
            for w in desc_words:
                if len(w) > 3 and w in dt:
                    score += 0.1
            if score > best_score:
                best_score = score
                best_id = tid
                best_desc = desc
                best_parent = pdesc or ''

    if best_id and best_score >= 0.18:
        return best_id, f'FUZZY (s={best_score:.2f}) {best_parent}>{best_desc}'

    return None, f'FAIL (best={best_score:.2f})'

# ============================================================
# 9. DESCRIPTION WASHER
# ============================================================
def wash_description(cur, item_template_id, raw_desc):
    cur.execute("""
        SELECT ct.category_name
        FROM item_template it
        JOIN category_template ct ON it.category_id = ct.id
        WHERE it.id = ?
    """, (item_template_id,))
    cat_row = cur.fetchone()
    cat_name = cat_row[0] if cat_row else 'UNKNOWN'

    # Tier 1: Item-specific
    cur.execute("""
        SELECT description FROM defect_library
        WHERE tenant_id = ? AND item_template_id = ?
        ORDER BY usage_count DESC
    """, (TENANT, item_template_id))
    item_entries = [r[0] for r in cur.fetchall()]
    if item_entries:
        best = None
        best_s = 0
        for entry in item_entries:
            s = fuzzy(raw_desc, entry)
            if s > best_s:
                best_s = s
                best = entry
        if best and best_s >= 0.7:
            return best, f"item (s={best_s:.2f})", cat_name

    # Tier 2: Category fallback
    cur.execute("""
        SELECT description FROM defect_library
        WHERE tenant_id = ? AND category_name = ? AND item_template_id IS NULL
        ORDER BY usage_count DESC
    """, (TENANT, cat_name))
    cat_entries = [r[0] for r in cur.fetchall()]
    if cat_entries:
        best = None
        best_s = 0
        for entry in cat_entries:
            s = fuzzy(raw_desc, entry)
            if s > best_s:
                best_s = s
                best = entry
        if best and best_s >= 0.7:
            return best, f"cat (s={best_s:.2f})", cat_name

    # Tier 3: Clean raw text
    cleaned = raw_desc.strip()
    if cleaned:
        cleaned = cleaned[0].upper() + cleaned[1:]
    return cleaned, "NEW", cat_name

# ============================================================
# 10. MAIN IMPORT
# ============================================================
def main():
    if len(sys.argv) < 3:
        print("Usage:")
        print("  python3 universal_import.py <docx_or_json> <cycle_id> [--commit]")
        print("  python3 universal_import.py --inline <cycle_id> [--commit] << 'DATA'")
        print("    unit|inspector_id|inspector_name|date")
        print("    AREA|sub_header|defect text")
        print("    ...")
        print("    DATA")
        sys.exit(1)

    commit = '--commit' in sys.argv
    inline_mode = '--inline' in sys.argv

    # --- PARSE ---
    print("=" * 60)
    print("PARSING INPUT")
    print("=" * 60)

    if inline_mode:
        cycle_id = sys.argv[sys.argv.index('--inline') + 1]
        # Check if next arg after cycle_id is a file path
        ci = sys.argv.index('--inline')
        txt_file = sys.argv[ci + 2] if len(sys.argv) > ci + 2 and not sys.argv[ci + 2].startswith('-') else None
        if txt_file and os.path.exists(txt_file):
            with open(txt_file) as f:
                lines = [l.strip() for l in f if l.strip()]
        else:
            lines = [l.strip() for l in sys.stdin if l.strip()]
        if not lines:
            print("ERROR: No data on stdin")
            sys.exit(1)
        # First line: unit|inspector_id|inspector_name|date
        meta = lines[0].split('|')
        unit_number = meta[0].strip()
        inspector_id = meta[1].strip()
        inspector_name = meta[2].strip()
        inspection_date = meta[3].strip()
        raw_defects = []
        for i, line in enumerate(lines[1:], 1):
            parts = line.split('|', 2)
            if len(parts) == 3:
                raw_defects.append((parts[0].strip(), parts[1].strip(),
                                   parts[2].strip(), i))
        print(f"  Loaded inline: {len(raw_defects)} defects")
    else:
        docx_path = sys.argv[1]
        cycle_id = sys.argv[2]
        if not os.path.exists(docx_path):
            print(f"ERROR: File not found: {docx_path}")
            sys.exit(1)
        if docx_path.endswith('.json'):
            import json
            with open(docx_path) as f:
                data = json.load(f)
            unit_number = data['unit_number']
            inspector_name = data['inspector_name']
            inspector_id = data['inspector_id']
            inspection_date = data['inspection_date']
            raw_defects = [(d['area'], d['sub_header'], d['text'], d.get('line', 0))
                           for d in data['defects']]
            print(f"  Loaded from JSON: {len(raw_defects)} defects")
        else:
            unit_number, inspector_name, inspection_date, raw_defects = parse_word_doc(docx_path)
            inspector_id = get_inspector_id(inspector_name) if inspector_name else None

    print(f"  Unit: {unit_number}")
    print(f"  Inspector: {inspector_name} ({inspector_id})")
    print(f"  Date: {inspection_date}")
    print(f"  Raw defects parsed: {len(raw_defects)}")

    if not unit_number or not inspector_id or not inspection_date:
        print("ERROR: Could not extract all metadata from doc")
        if not inspector_id and inspector_name:
            print(f"  Unknown inspector: '{inspector_name}' - add to INSPECTOR_MAP")
        sys.exit(1)

    # --- FILTER DROPS ---
    print()
    print("=" * 60)
    print("FILTERING DROPS")
    print("=" * 60)
    filtered = []
    for area, sub_header, text, line_num in raw_defects:
        sh_lower = sub_header.lower().strip()
        # Area+sub-header drop
        if (area, sh_lower) in AREA_SUB_DROPS:
            print(f"  DROP [area+sub] {area}>{sub_header}: {text}")
            continue
        # Content drop
        if should_drop_content(text):
            print(f"  DROP [content] {text}")
            continue
        filtered.append((area, sub_header, text, line_num))

    print(f"  Kept: {len(filtered)}, Dropped: {len(raw_defects) - len(filtered)}")

    # --- RESOLVE TEMPLATES ---
    print()
    print("=" * 60)
    print("TEMPLATE RESOLUTION")
    print("=" * 60)

    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()

    resolved = []
    failures = []

    for area, sub_header, text, line_num in filtered:
        category, parent_kw = get_sub_header_mapping(area, sub_header)

        if category is None:
            failures.append((area, sub_header, text, line_num,
                           f'UNKNOWN SUB-HEADER: {sub_header}'))
            print(f"  FAIL [unknown sub-header] {area}>{sub_header}: {text}")
            continue

        tmpl_id, match_info = resolve_template(cur, area, category, parent_kw, text)

        if tmpl_id:
            dtype = detect_defect_type(text)
            resolved.append((tmpl_id, text, dtype, area, category, parent_kw))
            status_char = 'O' if 'OVERRIDE' in match_info else 'F'
            print(f"  {status_char} [{tmpl_id}] {area}>{category}>{parent_kw}: {text}")
            print(f"    -> {match_info}")
        else:
            failures.append((area, sub_header, text, line_num, match_info))
            print(f"  FAIL {area}>{category}>{parent_kw}: {text}")
            print(f"    -> {match_info}")

    print()
    print(f"Resolved: {len(resolved)}, Failed: {len(failures)}")

    if failures:
        print()
        print("=" * 60)
        print(f"*** {len(failures)} UNRESOLVED - CANNOT IMPORT ***")
        print("=" * 60)
        for area, sub, text, ln, info in failures:
            print(f"  Line {ln}: {area} > {sub}: {text}")
            print(f"    {info}")
        print()
        print("Fix: Add overrides to TEMPLATE_OVERRIDES in universal_import.py")
        conn.close()
        sys.exit(1)

    # --- PRE-IMPORT CHECKS ---
    cur.execute('SELECT id FROM unit WHERE unit_number=? AND tenant_id=?',
                (unit_number, TENANT))
    row = cur.fetchone()
    if not row:
        print(f"ERROR: Unit {unit_number} not found in DB")
        conn.close()
        sys.exit(1)
    unit_id = row[0]

    # Check for existing inspection
    cur.execute('SELECT id, status FROM inspection WHERE unit_id=? AND cycle_id=?',
                (unit_id, cycle_id))
    row = cur.fetchone()
    existing_insp = row if row else None

    print()
    print("=" * 60)
    print("IMPORT SUMMARY")
    print("=" * 60)
    print(f"  Unit: {unit_number} (ID: {unit_id})")
    print(f"  Cycle: {cycle_id}")
    print(f"  Inspector: {inspector_name} ({inspector_id})")
    print(f"  Date: {inspection_date}")
    print(f"  Defects to import: {len(resolved)}")
    if existing_insp:
        print(f"  Existing inspection: {existing_insp[0]} (status={existing_insp[1]})")

    if not commit:
        print()
        print("*** DRY RUN - No changes made ***")
        print(f"*** Run with --commit to import ***")
        conn.close()
        return

    # --- IMPORT ---
    print()
    print("=" * 60)
    print("IMPORTING")
    print("=" * 60)
    now = now_iso()

    # Create or reuse inspection
    if existing_insp:
        insp_id = existing_insp[0]
        if existing_insp[1] not in ('not_started', 'in_progress'):
            print(f"ERROR: Inspection already at {existing_insp[1]}")
            conn.close()
            sys.exit(1)
        print(f"Reusing inspection: {insp_id}")
    else:
        insp_id = gen_id()
        cur.execute("""
            INSERT INTO inspection
            (id, tenant_id, unit_id, cycle_id, inspection_date,
             inspector_id, inspector_name, status, started_at, created_at, updated_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, 'in_progress', ?, ?, ?)
        """, (insp_id, TENANT, unit_id, cycle_id, inspection_date,
              inspector_id, inspector_name, now, now, now))
        print(f"Created inspection: {insp_id}")

    # Update inspector
    cur.execute("UPDATE inspection SET inspector_id=?, inspector_name=?, updated_at=? WHERE id=?",
                (inspector_id, inspector_name, now, insp_id))

    # Create inspection items (if not exists)
    cur.execute('SELECT COUNT(*) FROM inspection_item WHERE inspection_id=?', (insp_id,))
    existing_items = cur.fetchone()[0]
    if existing_items > 0:
        print(f"Inspection items exist: {existing_items}")
    else:
        cur.execute('SELECT id FROM item_template WHERE tenant_id=?', (TENANT,))
        templates = cur.fetchall()
        for t in templates:
            cur.execute("""
                INSERT INTO inspection_item
                (id, tenant_id, inspection_id, item_template_id, status, marked_at)
                VALUES (?, ?, ?, ?, 'pending', NULL)
            """, (gen_id(), TENANT, insp_id, t[0]))
        print(f"Created {len(templates)} inspection items")

    # Mark exclusions (copy from other inspections in same cycle)
    cur.execute("""
        SELECT DISTINCT ii.item_template_id
        FROM inspection_item ii
        JOIN inspection i ON ii.inspection_id = i.id
        WHERE i.cycle_id = ? AND ii.status = 'skipped' AND i.id != ?
    """, (cycle_id, insp_id))
    excluded_ids = set(r[0] for r in cur.fetchall())
    print(f"Exclusions from cycle: {len(excluded_ids)}")

    skipped = 0
    for eid in excluded_ids:
        cur.execute("""
            UPDATE inspection_item SET status='skipped', marked_at=?
            WHERE inspection_id=? AND item_template_id=? AND status='pending'
        """, (now, insp_id, eid))
        skipped += cur.rowcount
    print(f"Marked skipped: {skipped}")

    # Filter out defects on excluded items
    clean_defects = []
    dropped_excl = 0
    for tmpl_id, desc, dtype, area, cat, pkw in resolved:
        if tmpl_id in excluded_ids:
            print(f"  DROPPED (excluded): [{tmpl_id}] {desc}")
            dropped_excl += 1
        else:
            clean_defects.append((tmpl_id, desc, dtype))
    print(f"Exclusion overlap: dropped {dropped_excl}, clean {len(clean_defects)}")

    # Create defects
    new_library = []
    defect_count = 0
    for template_id, raw_desc, dtype in clean_defects:
        washed, wash_src, cat_name = wash_description(cur, template_id, raw_desc)
        if "NEW" in wash_src:
            new_library.append((template_id, cat_name, washed))

        defect_id = gen_id()
        defect_type = 'not_installed' if dtype == 'NI' else 'not_to_standard'
        cur.execute("""
            INSERT INTO defect
            (id, tenant_id, unit_id, item_template_id, raised_cycle_id,
             defect_type, status, original_comment, created_at, updated_at)
            VALUES (?, ?, ?, ?, ?, ?, 'open', ?, ?, ?)
        """, (defect_id, TENANT, unit_id, template_id, cycle_id,
              defect_type, washed, now, now))

        item_status = 'not_installed' if dtype == 'NI' else 'not_to_standard'
        cur.execute("""
            UPDATE inspection_item SET status=?, comment=?, marked_at=?
            WHERE inspection_id=? AND item_template_id=?
        """, (item_status, washed, now, insp_id, template_id))
        defect_count += 1

    print(f"Defects created: {defect_count}")

    # Mark remaining as OK
    cur.execute("""
        UPDATE inspection_item SET status='ok', marked_at=?
        WHERE inspection_id=? AND status='pending'
    """, (now, insp_id))
    print(f"Marked OK: {cur.rowcount}")

    # Add new library entries
    if new_library:
        for tid, cname, desc in new_library:
            cur.execute("""
                INSERT INTO defect_library
                (id, tenant_id, category_name, item_template_id, description,
                 usage_count, is_system, created_at)
                VALUES (?, ?, ?, ?, ?, 1, 0, ?)
            """, (gen_id(), TENANT, cname, tid, desc, now))
        print(f"New library entries: {len(new_library)}")

    # Set inspection status
    cur.execute("UPDATE inspection SET status='submitted', submitted_at=?, updated_at=? WHERE id=?",
                (now, now, insp_id))

    # --- VERIFY ---
    print()
    print("=" * 60)
    print("VERIFICATION")
    print("=" * 60)
    total = 0
    for status in ['skipped', 'ok', 'not_to_standard', 'not_installed', 'pending']:
        cur.execute('SELECT COUNT(*) FROM inspection_item WHERE inspection_id=? AND status=?',
                    (insp_id, status))
        c = cur.fetchone()[0]
        total += c
        print(f"  {status}: {c}")
    cur.execute('SELECT COUNT(*) FROM defect WHERE unit_id=? AND raised_cycle_id=? AND status=?',
                (unit_id, cycle_id, 'open'))
    print(f"  defects: {cur.fetchone()[0]}")
    print(f"  total items: {total} (expect 523)")

    pending_check = 0
    cur.execute('SELECT COUNT(*) FROM inspection_item WHERE inspection_id=? AND status=?',
                (insp_id, 'pending'))
    pending_check = cur.fetchone()[0]
    if pending_check > 0:
        print(f"  WARNING: {pending_check} items still pending!")

    conn.commit()
    print()
    print("COMMITTED SUCCESSFULLY")
    conn.close()

if __name__ == '__main__':
    main()
